"""Manifest-driven local document loading for fictional or approved knowledge files."""

import json
from pathlib import Path
from typing import Any

from docx import Document as WordDocument
from pydantic import BaseModel, Field, ValidationError
from pypdf import PdfReader

from app.rag.errors import KnowledgeDocumentError
from app.rag.ingestion.contracts import LoadedKnowledgeDocument
from app.rag.schemas import PolicyDocumentMetadata


class ManifestEntry(PolicyDocumentMetadata):
    file_path: str
    attributes: dict[str, Any] = Field(default_factory=dict)


class ManifestFile(BaseModel):
    documents: list[ManifestEntry] = Field(default_factory=list)


class ManifestDiscoveryResult(BaseModel):
    entries: list[tuple[Path, ManifestEntry]] = Field(default_factory=list)
    warnings: list[str] = Field(default_factory=list)

    model_config = {"arbitrary_types_allowed": True}


def discover_manifests(policy_root: Path) -> ManifestDiscoveryResult:
    entries: list[tuple[Path, ManifestEntry]] = []
    warnings: list[str] = []
    if not policy_root.exists():
        return ManifestDiscoveryResult(warnings=["POLICY_DATA_DIR_NOT_FOUND"])
    for manifest_path in sorted(policy_root.rglob("manifest.json")):
        try:
            payload = json.loads(manifest_path.read_text(encoding="utf-8"))
            manifest = ManifestFile.model_validate(payload)
        except (OSError, json.JSONDecodeError, ValidationError):
            warnings.append(f"MANIFEST_INVALID:{manifest_path.name}")
            continue
        for entry in manifest.documents:
            source_path = (manifest_path.parent / entry.file_path).resolve()
            try:
                source_path.relative_to(policy_root.resolve())
            except ValueError:
                warnings.append(f"DOCUMENT_PATH_OUTSIDE_POLICY_ROOT:{entry.source_id}")
                continue
            entries.append((source_path, entry))
    if not entries and not warnings:
        warnings.append("NO_MANIFEST_DOCUMENTS")
    return ManifestDiscoveryResult(entries=entries, warnings=warnings)


class LocalKnowledgeLoader:
    SUPPORTED_SUFFIXES = {".md", ".txt", ".pdf", ".docx"}

    def load(self, source_path: str, metadata: ManifestEntry) -> LoadedKnowledgeDocument:
        path = Path(source_path)
        suffix = path.suffix.casefold()
        if suffix not in self.SUPPORTED_SUFFIXES:
            raise KnowledgeDocumentError("知识文件格式不受支持。")
        try:
            if suffix in {".md", ".txt"}:
                content = path.read_text(encoding="utf-8")
            elif suffix == ".pdf":
                pages = [page.extract_text() or "" for page in PdfReader(str(path)).pages]
                content = "\n\n".join(page.strip() for page in pages if page.strip())
                if not content:
                    raise KnowledgeDocumentError("PDF_NO_EXTRACTABLE_TEXT")
            else:
                document = WordDocument(str(path))
                content = "\n".join(paragraph.text for paragraph in document.paragraphs if paragraph.text.strip())
        except KnowledgeDocumentError:
            raise
        except (OSError, ValueError, RuntimeError, KeyError) as exc:
            raise KnowledgeDocumentError("知识文件读取失败。") from exc
        normalized = content.replace("\x00", "").strip()
        if not normalized:
            raise KnowledgeDocumentError("知识文件没有可索引文本。")
        base_metadata = PolicyDocumentMetadata.model_validate(
            metadata.model_dump(exclude={"file_path", "attributes"})
        )
        return LoadedKnowledgeDocument(
            metadata=base_metadata,
            content=normalized,
            attributes=metadata.attributes,
        )
