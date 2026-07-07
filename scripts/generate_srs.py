"""
生成 TalentFlow 智聘中枢《软件需求规格说明书》docx
"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ── 全局样式 ──
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(4)

# 标题样式
for lvl in range(1, 4):
    hs = doc.styles[f'Heading {lvl}']
    hf = hs.font
    hf.name = '黑体'
    hf.color.rgb = RGBColor(0, 0, 0)
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    if lvl == 1:
        hf.size = Pt(16)
        hf.bold = True
    elif lvl == 2:
        hf.size = Pt(14)
        hf.bold = True
    else:
        hf.size = Pt(13)
        hf.bold = True

# ── 封面 ──
for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('武汉大学计算机学院')
run.font.size = Pt(18)
run.font.name = '黑体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for ch in '软件需求规格说明书':
    run = p.add_run(ch + '\n')
    run.font.size = Pt(26)
    run.font.name = '黑体'
    run.font.bold = True
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

for _ in range(3):
    doc.add_paragraph()

cover_info = [
    ('组序号：', '第15组'),
    ('项目名称：', 'TalentFlow 智聘中枢'),
    ('专业（班）：', '软件工程'),
]
for label, value in cover_info:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'{label}{value}')
    run.font.size = Pt(14)

doc.add_paragraph()

members = [
    ('黄钧', '2024302131066', 'PO / 组长'),
    ('吴越', '2024302181013', 'SM / 设计负责人'),
    ('刘子恒', '2024302111006', 'QA / 编码负责人'),
    ('唐丞杰', '2024302181049', 'QA / 前端负责人'),
]
table = doc.add_table(rows=1 + len(members), cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = 'Table Grid'
for i, h in enumerate(['姓名', '学号', '角色']):
    cell = table.rows[0].cells[i]
    cell.text = h
    for par in cell.paragraphs:
        par.alignment = WD_ALIGN_PARAGRAPH.CENTER
for idx, (name, sid, role) in enumerate(members):
    for j, val in enumerate([name, sid, role]):
        cell = table.rows[idx + 1].cells[j]
        cell.text = val
        for par in cell.paragraphs:
            par.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('任课教师：杜卓敏')
run.font.size = Pt(14)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('2026 年 7 月 7 日')
run.font.size = Pt(14)

doc.add_page_break()

# ── 目录页 ──
doc.add_heading('目录', level=1)
toc = [
    ('1  引言', 3),
    ('   1.1  编写目的', 3),
    ('   1.2  读者对象', 3),
    ('   1.3  软件项目概述', 3),
    ('   1.4  文档概述', 4),
    ('   1.5  术语与定义', 4),
    ('   1.6  参考资料', 4),
    ('2  软件的一般性描述', 5),
    ('   2.1  系统架构与环境', 5),
    ('   2.2  任务目标', 5),
    ('   2.3  用户的特点', 6),
    ('   2.4  假定和约束', 6),
    ('3  软件功能需求描述', 7),
    ('   3.1  总体功能概述', 7),
    ('   3.2  核心用例描述', 7),
    ('4  需求规定', 8),
    ('   4.1  对功能的规定', 8),
    ('   4.2  对性能的规定', 9),
    ('   4.3  输入输出要求', 9),
    ('   4.4  故障处理要求', 9),
    ('   4.5  安全与权限要求', 10),
    ('5  数据要求说明', 10),
    ('6  运行环境规定', 11),
    ('7  验收与评审要求', 12),
]
for text, page in toc:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.8
    run = p.add_run(f'{text}{"." * (50 - len(text))}{page}')
    run.font.size = Pt(12)

doc.add_page_break()

# ═══════════════════════════════════════════
# 第1章 引言
# ═══════════════════════════════════════════
doc.add_heading('1  引言', level=1)

doc.add_heading('1.1  编写目的', level=2)
doc.add_paragraph(
    '本说明书用于明确"TalentFlow 智聘中枢"系统在功能、数据、性能、接口、运行环境和验收方面的需求。'
    '项目组后续进行概要设计、详细设计、编码实现、接口联调和验收测试时，应以其中确定的功能边界、'
    '数据约束和质量要求为依据。'
)
doc.add_paragraph(
    'TalentFlow 的需求重点是将招聘决策、员工服务、薪资权限与审计追踪组织成完整业务闭环。系统既要支持 '
    'HR 高效筛选候选人并安排面试，也要保证每一条薪资查询都能被权限校验和审计日志覆盖；既要利用 AI Agent '
    '串联跨模块任务，又要保证核心评分、排期和权限算法完全可解释、可追溯。'
)

doc.add_heading('1.2  读者对象', level=2)
doc.add_paragraph(
    '本说明书的读者包括项目指导教师、课程评审人员、项目组全体成员。需求分析人员应重点关注系统边界、参与者、'
    '核心用例和数据流；设计与编码人员应重点关注功能规定、数据要求、接口约束和AI禁飞区边界；测试人员应重点'
    '关注性能指标、异常流程和验收条件。'
)

doc.add_heading('1.3  软件项目概述', level=2)
items = [
    '项目名称：TalentFlow 智聘中枢。',
    '项目类型：面向招聘决策、员工服务与权限审计的可解释企业人力资源管理 Agent 系统。',
    '任务提出者：武汉大学计算机学院软件工程课程项目组。',
    '开发单位：第15组（四qu战车）项目组。',
    '最终用户：HR 专员、部门主管、薪酬管理员和普通员工。',
    '运行环境：采用本地 Docker Compose + Nginx 部署演示环境。',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph(
    'TalentFlow 面向的是企业 HR 工作中"岗位需求—候选人筛选—面试排期—薪资权限—审计追溯"的连续管理需求。'
    '系统定位不是"普通 HR 系统加聊天框"，而是让 Agent 参与核心人力资源业务流程。'
)

doc.add_heading('1.4  文档概述', level=2)
doc.add_paragraph(
    '全文按软件需求、数据要求和运行环境三条主线组织。第1章说明背景与范围；第2章说明系统架构、任务目标和约束条件；'
    '第3章围绕核心用例描述系统功能；第4章给出功能、性能、安全等方面的需求规定；第5章说明数据分类与存储要求；'
    '第6章说明设备、支持软件和接口要求；第7章给出验收与评审要求。'
)

doc.add_heading('1.5  术语与定义', level=2)
terms = [
    ('TalentFlow / 智聘中枢', '本项目拟开发的企业人力资源管理 Agent 系统。'),
    ('AI 禁飞区', '必须由人工手写的三个核心算法模块，AI 不得生成或修改其实现代码。'),
    ('LangGraph', '用于编排员工服务 Agent 与招聘决策 Agent 执行流程的框架。'),
    ('RAG', 'Retrieval-Augmented Generation，检索增强生成。系统通过 ChromaDB 检索企业制度后由 LLM 生成带来源的回答。'),
    ('Gradio', '内部 Agent 调试台，用于查看 LangGraph 执行链、工具调用和 RAG 命中内容，不作为正式业务入口。'),
    ('SSE', 'Server-Sent Events，用于前端接收 Agent 流式响应的推送机制。'),
]
table = doc.add_table(rows=1 + len(terms), cols=2)
table.style = 'Table Grid'
table.rows[0].cells[0].text = '术语'
table.rows[0].cells[1].text = '定义'
for i, (term, defn) in enumerate(terms):
    table.rows[i + 1].cells[0].text = term
    table.rows[i + 1].cells[1].text = defn

doc.add_heading('1.6  参考资料', level=2)
refs = [
    '《计算机软件产品开发文件编制指南：软件需求说明书（GB856T–88）》；',
    '《计算机软件产品开发文件编制指南：数据要求说明书（GB856T–88）》；',
    '《TalentFlow 智聘中枢计划任务书》（docs/计划.md）；',
    '《TalentFlow 智聘中枢架构说明》（.agent/architecture.md）；',
    'Vue 3、FastAPI、LangGraph、PostgreSQL、ChromaDB 相关技术文档。',
]
for r in refs:
    doc.add_paragraph(r, style='List Number')

doc.add_page_break()

# ═══════════════════════════════════════════
# 第2章 软件的一般性描述
# ═══════════════════════════════════════════
doc.add_heading('2  软件的一般性描述', level=1)

doc.add_heading('2.1  系统架构与环境', level=2)
doc.add_paragraph(
    '本软件采用前后端分离的 B/S 架构。Vue 3 企业工作台负责用户交互、图表展示和 Agent 过程可视化；'
    'FastAPI 后端负责身份认证、业务编排、Agent 任务调度和数据持久化；PostgreSQL 保存结构化业务数据；'
    'ChromaDB 保存企业制度知识库向量并支持 RAG 检索；LangGraph Agent 编排员工服务与招聘决策等跨模块任务；'
    'Gradio 仅用于内部 Agent 调试。系统整体通过 Docker Compose + Nginx 进行本地部署。'
)
doc.add_paragraph(
    '系统的外部环境可以理解为三层：用户侧浏览器环境、服务端运行环境和 AI 分析环境。浏览器环境直接决定 '
    '图表渲染、日历交互和 SSE 流式接收的可用性；服务端环境决定接口稳定性、数据库事务和 Agent 任务调度能力；'
    'AI 分析环境决定 RAG 问答质量、候选人评分解释和排期推荐的可信度。'
)
doc.add_paragraph(
    '普通页面操作走"Vue → FastAPI → Service → Repository / Database"链路；自然语言、多工具和跨模块任务走'
    '"Vue → FastAPI → LangGraph Agent → Tool → Service"链路；制度问答通过 ChromaDB 检索并由 LLM 生成带来源的回答。'
)

doc.add_heading('2.2  任务目标', level=2)
goals = [
    '构建企业级 SaaS 工作台，支持多角色登录（普通员工、HR 专员、部门主管、薪酬管理员）；',
    '实现员工自助服务 Agent，支持假期、本人薪资和公司制度查询，展示制度依据和来源片段；',
    '实现智能招聘决策，辅助岗位画像、候选人筛选、多维加权评分、权重沙盘和候选人多维对比；',
    '实现智能面试排期，结合候选人、面试官、会议室和时间槽生成推荐方案并说明冲突原因；',
    '实现薪资权限控制，按角色和字段级别控制薪资可见范围，记录所有敏感访问审计日志；',
    '实现招聘驾驶舱与报告，展示招聘漏斗、技能缺口、面试官负载和排期风险；',
    '通过 Gradio 内部调试台查看 Agent 执行链、工具调用、RAG 命中内容和错误信息。',
]
for g in goals:
    doc.add_paragraph(g, style='List Number')

doc.add_paragraph(
    '这些目标之间存在依赖关系。没有稳定的身份认证和角色管理，薪资权限就无从校验；没有结构化的候选人数据和'
    '评分算法，智能筛选就无法提供可解释的排序；没有制度知识库和 RAG 检索，员工问答就只能给出无来源的结论。'
    '因此项目的需求优先级以"评分—排期—权限"三个禁飞区为中心，向外围 Agent 编排、RAG 和前端可视化逐步扩展。'
)

doc.add_heading('2.3  用户的特点', level=2)
doc.add_paragraph(
    '系统面向四类用户。普通员工关注假期、本人薪资和制度查询，他们通常通过 Agent 自然语言提问获取信息，'
    '不愿意在多个页面之间跳转查找。HR 专员是最活跃的用户群体，他们需要创建岗位、筛选候选人、调整评分权重、'
    '查看多维对比、管理招聘流程看板和安排面试。部门主管关注招聘进度、团队技能缺口和高层次分析报告。'
    '薪酬管理员需要控制薪资可见范围、查看敏感访问审计日志并确保权限配置正确。'
)
doc.add_paragraph(
    '不同用户的共同特点是对"可解释性"有明确需求。候选人排序需要展示评分维度和权重，薪资查询需要展示字段级'
    '脱敏效果，面试排期需要说明冲突原因，制度问答需要标注来源片段。系统不能把核心判断包装成不可理解的黑盒。'
)

doc.add_heading('2.4  假定和约束', level=2)
constraints = [
    '项目周期为 12 天（3 个 Sprint），团队 4 人，采用模块化单体架构而非微服务；',
    '正式演示目标为本地 Docker Compose + Nginx 部署，不引入 Kubernetes 或云服务；',
    'Agent 只处理自然语言、多工具、跨模块任务，普通 CRUD 操作走传统业务 API；',
    'AI 禁飞区（简历评分、面试排期、薪资权限）必须由人工手写，AI 不得生成或修改其实现代码；',
    'LLM 接口使用 OpenAI 兼容模型，需预留模型不可用时的降级方案；',
    '核心评分、排期、权限模块即使模型接口临时异常也应能独立运行和展示。',
]
for c in constraints:
    doc.add_paragraph(c, style='List Bullet')

doc.add_page_break()

# ═══════════════════════════════════════════
# 第3章 软件功能需求描述
# ═══════════════════════════════════════════
doc.add_heading('3  软件功能需求描述', level=1)

doc.add_heading('3.1  总体功能概述', level=2)
doc.add_paragraph(
    'TalentFlow 的功能需求围绕三类核心参与者展开。HR 端覆盖岗位创建、岗位画像、简历导入、候选人筛选、'
    '权重沙盘、多维对比、流程看板、面试排期和招聘驾驶舱。员工端覆盖多角色登录、假期查询、本人薪资查询、'
    '制度问答和 Agent 交互。管理端覆盖制度知识库、薪资权限控制、敏感访问审计和通知提醒。'
)
doc.add_paragraph(
    '功能可概括为三组。第一组是基础能力：身份认证、角色管理、权限上下文、统一响应和 Trace ID 追踪。'
    '第二组是业务能力：岗位管理、候选人评分、权重沙盘、多维对比、招聘流程、面试排期、员工服务和薪资查询。'
    '第三组是 Agent 与智能能力：员工服务 Agent、招聘决策 Agent、LangGraph 编排、Agent Tools、RAG 制度检索和 Gradio 调试台。'
)

doc.add_heading('3.2  核心用例描述', level=2)

# --- 用例1 ---
doc.add_heading('3.2.1  候选人智能筛选与评分', level=3)
doc.add_paragraph(
    'HR 创建岗位并输入岗位描述后，系统应输出结构化岗位要求。简历导入后系统通过简历评分算法对候选人进行'
    '多维加权评分，评分维度至少包括技能匹配、项目经历和到岗时间。HR 可以在权重沙盘中调整各维度权重，'
    '候选人排名和评分解释实时变化。HR 还可进行候选人多维对比，查看技能、项目、综合评分等维度的差距。'
)
doc.add_paragraph('• 输入：岗位描述、候选人简历、评分权重配置。')
doc.add_paragraph('• 输出：候选人排序列表、分维度评分、评分依据解释、多维对比视图。')
doc.add_paragraph('• 约束：评分至少包含三个维度；权重改变后排序应实时变化；评分结果需可解释每个维度的得分依据。')

# --- 用例2 ---
doc.add_heading('3.2.2  智能面试排期', level=3)
doc.add_paragraph(
    'HR 选择候选人后，系统结合候选人空闲时段、面试官可用时间、会议室资源和时间槽，通过排期算法生成推荐面试方案。'
    '排期结果应避免候选人、面试官和会议室三类冲突。无法排期时应返回具体原因（如"候选人该时段不可用"或"无可用会议室"）。'
    '系统还应展示备选时段和冲突说明。'
)
doc.add_paragraph('• 输入：候选人列表、面试官可用时间、会议室资源、时间槽约束。')
doc.add_paragraph('• 输出：推荐排期方案、备选时段、冲突原因说明。')
doc.add_paragraph('• 约束：排期必须避免三类冲突；无法排期时必须返回可理解的原因。')

# --- 用例3 ---
doc.add_heading('3.2.3  员工自助服务与制度问答', level=3)
doc.add_paragraph(
    '员工通过底部 Agent 指令栏或员工端页面进行自然语言提问，如"我还有多少年假""下周请三天是否符合制度"等。'
    '系统通过 LangGraph 员工服务 Agent 编排流程：查询假期余额、检索相关制度条款、判断是否符合政策规定，'
    '最终返回答案并展示制度依据和来源片段。员工也可查询本人薪资信息。'
)
doc.add_paragraph('• 输入：自然语言问题（假期、薪资、制度相关）。')
doc.add_paragraph('• 输出：答案、制度来源片段、工具调用过程展示。')
doc.add_paragraph('• 约束：制度问答必须返回来源依据，不能只给结论。')

# --- 用例4 ---
doc.add_heading('3.2.4  薪资权限控制与审计', level=3)
doc.add_paragraph(
    '薪资查询按照角色、访问对象和字段级别进行权限校验。普通员工只能查看本人薪资完整信息；'
    'HR 查看他人薪资时仅获得脱敏范围（如薪资等级或区间）；薪酬管理员可查看完整薪资字段。'
    '每次薪资查询（无论允许或拒绝）均写入审计日志，包含访问者、访问时间、访问对象、查看字段和权限结果。'
)
doc.add_paragraph('• 输入：查询者角色、被查对象、请求字段。')
doc.add_paragraph('• 输出：按权限返回完整/脱敏/拒绝结果，审计日志记录。')
doc.add_paragraph('• 约束：同一条请求在不同角色下必须返回不同数据范围；所有访问必须有审计记录。')

# --- 用例5 ---
doc.add_heading('3.2.5  招聘驾驶舱与报告', level=3)
doc.add_paragraph(
    'HR 可在招聘驾驶舱中查看招聘漏斗（各阶段候选人数量分布）、技能缺口分析、面试官负载分布和排期风险评估。'
    '系统通过 ECharts 图表展示关键指标，支持导出招聘报告。驾驶舱数据与候选人池、面试排期、流程看板实时联动。'
)
doc.add_paragraph('• 输入：招聘流程数据、候选人数据、面试排期数据。')
doc.add_paragraph('• 输出：招聘漏斗图、技能缺口雷达图、面试官负载图、排期风险提示、可导出的招聘报告。')

# --- 用例6 ---
doc.add_heading('3.2.6  Agent 过程可视化与调试', level=3)
doc.add_paragraph(
    '用户在 Vue 工作台中可查看 Agent 的真实工具调用步骤，而不是只看到一句最终回答。前端通过 SSE 接收 Agent 流式响应，'
    '展示 LangGraph 执行链（Input → Tool → LLM 节点）、工具调用时间线和 RAG 命中来源抽屉。'
    '开发人员可通过 Gradio 内部调试台查看更详细的 Agent 执行过程、检索结果和错误信息。'
)
doc.add_paragraph('• 输入：用户 Agent 指令。')
doc.add_paragraph('• 输出：工具调用时间线、RAG 来源抽屉、LangGraph 节点状态、错误信息。')
doc.add_paragraph('• 约束：用户应能看到 Agent 执行过程而不仅是最终结论；Gradio 不作为正式业务入口。')

doc.add_page_break()

# ═══════════════════════════════════════════
# 第4章 需求规定
# ═══════════════════════════════════════════
doc.add_heading('4  需求规定', level=1)

doc.add_heading('4.1  对功能的规定', level=2)
doc.add_paragraph(
    'TalentFlow 系统的功能需求应覆盖从身份认证到 Agent 智能决策的完整业务链路。各功能之间不是孤立页面，'
    '而是围绕 HR 招聘管理和员工自助服务形成连续闭环。'
)

# 功能需求摘要表
funcs = [
    ('FR-01', '多角色登录与认证', '支持普通员工、HR 专员、部门主管、薪酬管理员四种角色，基于 Token 鉴权。'),
    ('FR-02', '岗位创建与岗位画像', 'HR 输入岗位描述，系统输出结构化岗位要求。'),
    ('FR-03', '简历导入与智能筛选', '支持简历导入，通过多维加权评分对候选人排序并给出可解释的评分依据。'),
    ('FR-04', '招聘权重沙盘', 'HR 可调整技能、项目、到岗时间等权重，候选人排名和解释实时变化。'),
    ('FR-05', '候选人多维对比', '对比候选人在技能、项目、到岗时间、综合评分等维度的差异。'),
    ('FR-06', '招聘流程看板', '管理投递、初筛、约面、面试、Offer、入职或淘汰等阶段的候选人流转。'),
    ('FR-07', '智能面试排期', '结合候选人、面试官、会议室和时间槽生成推荐方案并说明冲突原因。'),
    ('FR-08', '员工 Agent 自助服务', '员工通过自然语言查询假期、本人薪资和公司制度，返回制度来源依据。'),
    ('FR-09', '薪资权限控制', '按角色和字段级别控制薪资可见范围，支持完整、脱敏和拒绝三种状态。'),
    ('FR-10', '敏感访问审计', '记录薪资查询、候选人评分和 Agent 操作等关键访问日志，支持按时间、角色、结果筛选。'),
    ('FR-11', '制度知识库 RAG', '企业制度通过 ChromaDB 进行向量检索，回答问题时展示命中片段和制度来源名称。'),
    ('FR-12', '招聘驾驶舱与报告', '展示招聘漏斗、技能缺口、面试官负载、排期风险和 KPI 卡片，支持导出招聘报告。'),
    ('FR-13', 'Agent 过程可视化', '前端展示 Agent 工具调用时间线、LangGraph 执行节点、RAG 来源抽屉和错误信息。'),
    ('FR-14', 'Gradio 调试台', '内部调试查看 LangGraph 执行链、工具调用、RAG 检索结果和 Agent 错误信息。'),
]
table = doc.add_table(rows=1 + len(funcs), cols=3)
table.style = 'Table Grid'
table.rows[0].cells[0].text = '编号'
table.rows[0].cells[1].text = '功能名称'
table.rows[0].cells[2].text = '核心要求'
for i, (n, name, req) in enumerate(funcs):
    table.rows[i + 1].cells[0].text = n
    table.rows[i + 1].cells[1].text = name
    table.rows[i + 1].cells[2].text = req

doc.add_paragraph(
    '系统在 Sprint 1 应完成"岗位—筛选—排期—权限"的基础闭环；Sprint 2 完成 Agent 编排、RAG 检索和前端 Agent 过程可视化；'
    'Sprint 3 完成页面打磨、报告生成、Docker 部署和答辩准备。各模块验收时不仅要看页面是否存在，也要看数据是否进入正确表、'
    '权限是否正确执行、审计日志是否完整记录。'
)

doc.add_heading('4.2  对性能的规定', level=2)
doc.add_paragraph(
    '普通业务查询接口（岗位列表、候选人查询、员工信息等）在正常负载下响应时间应控制在 2 秒以内。'
    'Agent 流式接口应在 1 秒内返回首个 SSE 事件，使用户感知到系统已开始处理。'
    'LLM 调用允许更长耗时，但前端必须通过 SSE 持续接收中间状态（工具调用、节点切换等）。'
    '简历评分和面试排期算法应在 500ms 内完成计算并返回结果。'
)
doc.add_paragraph(
    '• 评分结果必须包含各维度得分和计算依据，支持"可解释、可变化、可追溯"。'
    '• 排期结果必须说明主要约束和冲突原因。'
    '• 权限判断结果必须记录判断依据（角色、访问对象、字段级规则）。'
)

doc.add_heading('4.3  输入输出要求', level=2)
io_table = [
    ('用户认证', '用户名、密码、角色选择', 'Token、用户信息、角色权限'),
    ('岗位管理', '岗位描述、岗位要求', '结构化岗位画像'),
    ('候选人评分', '简历数据、权重配置', '排序列表、分维度评分、评分依据'),
    ('面试排期', '候选人、面试官、会议室、时间槽', '推荐方案、备选时段、冲突说明'),
    ('薪资查询', '查询者角色、被查对象ID、请求字段', '完整/脱敏/拒绝结果'),
    ('制度问答', '自然语言问题', '答案、来源片段、来源名称'),
    ('审计查询', '筛选条件（时间、角色、结果）', '审计日志列表、访问详情'),
    ('招聘报告', '报告参数（时间范围、岗位范围）', '招聘漏斗图、技能缺口分析、排期风险'),
]
table = doc.add_table(rows=1 + len(io_table), cols=3)
table.style = 'Table Grid'
table.rows[0].cells[0].text = '功能类别'
table.rows[0].cells[1].text = '典型输入'
table.rows[0].cells[2].text = '典型输出'
for i, (cat, inp, outp) in enumerate(io_table):
    table.rows[i + 1].cells[0].text = cat
    table.rows[i + 1].cells[1].text = inp
    table.rows[i + 1].cells[2].text = outp

doc.add_heading('4.4  故障处理要求', level=2)
faults = [
    'LLM 接口不可用时，核心评分、排期和权限模块仍应保持可用，Agent 功能可降级为提示"AI 服务暂不可用"；',
    '数据库连接异常时，后端应返回明确的错误提示而非崩溃，前端应展示友好的错误状态页；',
    '权限判断失败（如数据库查询异常）时，默认拒绝访问并记录异常审计日志，不允许因技术故障而绕过权限；',
    'RAG 检索无命中结果时，Agent 应如实告知"未找到相关制度"，不得编造内容；',
    'SSE 连接中断时，前端应展示断连提示并提供重试入口；',
    '所有失败请求应通过 Trace ID 可追踪。',
]
for f in faults:
    doc.add_paragraph(f, style='List Bullet')

doc.add_heading('4.5  安全与权限要求', level=2)
doc.add_paragraph(
    '用户密码必须哈希保存，受保护接口必须进行 Token 鉴权。薪资查询、候选人评分和 Agent 敏感操作必须经过权限校验。'
    '所有敏感访问（无论允许或拒绝）必须写入审计日志，包含访问者、时间、对象、字段和结果。'
    'AI 禁飞区代码不得由 AI 生成或修改，Agent 调用核心算法必须遵循"Tool → Service → human_only"链路，'
    '不得绕过 Service 层权限检查或直接访问禁飞区内部实现。真实 API Key、密码、Token 不得提交到代码仓库。'
)

doc.add_page_break()

# ═══════════════════════════════════════════
# 第5章 数据要求说明
# ═══════════════════════════════════════════
doc.add_heading('5  数据要求说明', level=1)

doc.add_heading('5.1  数据的逻辑描述', level=2)
doc.add_paragraph(
    'TalentFlow 的数据以用户和组织为中心组织，围绕"用户—岗位—候选人—面试—薪资—审计—制度"形成逻辑链。'
    '用户（员工）通过角色获得权限，岗位由HR创建并关联候选人，候选人参与招聘流程并关联评分和排期结果，'
    '薪资数据按角色和字段级别控制可见性，审计日志记录所有敏感访问，企业制度以向量形式存储在 ChromaDB 中。'
)
doc.add_paragraph(
    '数据分为三类：静态数据（系统配置、角色字典、制度知识库）、动态输入数据（用户操作、岗位创建、简历导入、'
    '评分请求、排期请求、薪资查询）、动态输出数据（评分结果、排期方案、权限判断、审计日志、招聘报告）。'
)

doc.add_heading('5.2  核心数据实体', level=2)
entities = [
    '用户（User）与角色（Role）：用户是数据归属中心，角色决定权限边界。',
    '岗位（Job）：包含岗位描述、岗位画像、要求技能、状态等。',
    '候选人（Candidate）：包含简历信息、技能标签、项目经历、到岗时间、评分结果等。',
    '面试排期（Interview Schedule）：关联候选人、面试官、会议室、时间槽和排期结果。',
    '薪资记录（Salary）：关联员工，包含基本薪资、津贴、奖金等字段，按角色控制可见范围。',
    '审计日志（Audit Log）：记录访问者、时间、操作类型、目标对象、访问字段和权限结果。',
    '制度知识库（Policy Knowledge Base）：企业制度文档通过 ChromaDB 进行向量化存储和检索。',
    '通知（Notification）：流程状态变更、面试安排、待办事项等站内通知。',
]
for e in entities:
    doc.add_paragraph(e, style='List Bullet')

doc.add_heading('5.3  数据约定', level=2)
doc.add_paragraph(
    '所有业务数据必须关联用户或具备明确归属关系。时间字段统一采用 ISO 8601 格式。薪资相关字段需在数据库层'
    '和 API 层双重控制可见性，不得仅依赖前端隐藏。审计日志数据只增不删，保留完整的敏感操作溯源链。'
    '制度知识库文档变更时需触发 ChromaDB 向量索引更新。密码不得明文存储，Token 需设置合理的有效期。'
)

doc.add_page_break()

# ═══════════════════════════════════════════
# 第6章 运行环境规定
# ═══════════════════════════════════════════
doc.add_heading('6  运行环境规定', level=1)

doc.add_heading('6.1  设备', level=2)
env_table = [
    ('客户端设备', '支持现代浏览器的 PC，建议 Chrome/Edge 最新版，分辨率 ≥ 1366×768。'),
    ('应用服务器', '项目组笔记本电脑，通过 Docker 容器提供后端服务。'),
    ('数据库服务器', 'Docker 化 PostgreSQL 16 实例，SSD 存储。'),
    ('向量数据库', 'Docker 化 ChromaDB 实例，与后端部署在同一网络。'),
]
table = doc.add_table(rows=1 + len(env_table), cols=2)
table.style = 'Table Grid'
table.rows[0].cells[0].text = '设备类别'
table.rows[0].cells[1].text = '要求'
for i, (cat, req) in enumerate(env_table):
    table.rows[i + 1].cells[0].text = cat
    table.rows[i + 1].cells[1].text = req

doc.add_heading('6.2  支持软件', level=2)
sw_table = [
    ('操作系统', 'Windows 11 / Ubuntu 22.04 LTS', '服务端运行基础。'),
    ('前端框架', 'Vue 3 + TypeScript + Vite + Element Plus + ECharts + FullCalendar', '构建企业 SaaS 工作台。'),
    ('后端框架', 'Python 3.12 + FastAPI + Uvicorn', '提供异步 Web API 服务。'),
    ('数据库', 'PostgreSQL 16', '持久保存结构化和业务数据。'),
    ('向量数据库', 'ChromaDB', '存储和检索企业制度知识库向量。'),
    ('Agent 框架', 'LangGraph + LangChain', '编排 Agent 执行流程和工具调用。'),
    ('调试工具', 'Gradio', '内部 Agent 调试台，查看执行链和工具调用。'),
    ('部署工具', 'Docker + Docker Compose + Nginx', '容器化部署和反向代理。'),
    ('接口文档', 'Swagger / OpenAPI（FastAPI 自带）', '前后端接口契约与调试。'),
    ('测试工具', 'Pytest', '后端单元测试与集成测试。'),
]
table = doc.add_table(rows=1 + len(sw_table), cols=3)
table.style = 'Table Grid'
table.rows[0].cells[0].text = '软件类别'
table.rows[0].cells[1].text = '软件名称'
table.rows[0].cells[2].text = '说明'
for i, (cat, name, desc) in enumerate(sw_table):
    table.rows[i + 1].cells[0].text = cat
    table.rows[i + 1].cells[1].text = name
    table.rows[i + 1].cells[2].text = desc

doc.add_heading('6.3  接口', level=2)
interfaces = [
    '前后端接口：采用 RESTful API，数据格式为 JSON。Agent 流式响应使用 SSE（Server-Sent Events）。',
    '认证接口：采用 Token 鉴权（JWT），登录后由后端返回访问令牌。',
    '数据库接口：后端通过 SQLAlchemy ORM 访问 PostgreSQL。',
    '向量检索接口：后端通过 ChromaDB 客户端进行制度知识库检索。',
    'AI 接口：LLM 通过 OpenAI 兼容接口调用，通过适配层封装以支持超时、重试和降级。',
    '调试接口：开发阶段保留 Swagger 文档（/docs）和 Gradio 调试台入口。',
]
for iface in interfaces:
    doc.add_paragraph(iface, style='List Bullet')

doc.add_heading('6.4  控制', level=2)
doc.add_paragraph(
    '系统控制来自用户操作、后端业务规则和 Agent 编排三个层面。用户通过 Vue 工作台页面和 Agent 指令栏触发请求，'
    '后端通过 API 路由、权限中间件和 Service 层决定请求的执行方式。管理员（导师）可以通过 Gradio 调试台查看 '
    'Agent 执行状态，但不直接干预业务数据。控制策略应优先保护核心功能：当 LLM 服务不可用时，保留评分、排期和'
    '权限三大核心模块的独立运行能力。'
)

doc.add_page_break()

# ═══════════════════════════════════════════
# 第7章 验收与评审要求
# ═══════════════════════════════════════════
doc.add_heading('7  验收与评审要求', level=1)

doc.add_heading('7.1  验收要求', level=2)
doc.add_paragraph(
    '系统验收应围绕三个 Sprint 的核心交付成果进行。'
)
sprint1 = [
    'HR 能创建岗位、导入候选人、查看评分和排序结果；',
    '调整评分权重后候选人排名实时变化；',
    '排期方案能避免候选人、面试官、会议室三类冲突；',
    '普通员工查看本人薪资允许，查看他人薪资拒绝，HR 查看他人薪资获得脱敏范围；',
    '所有敏感访问均有审计记录；',
    'Vue 工作台页面视觉统一，各模块可切换。',
]
doc.add_paragraph('Sprint 1（核心业务闭环）：')
for s in sprint1:
    doc.add_paragraph(s, style='List Bullet')

sprint2 = [
    '员工 Agent 可完成"假期余额 + 制度判断"问答；',
    '招聘 Agent 可完成"筛选候选人 + 排期"任务；',
    'RAG 回答带制度来源片段和来源名称；',
    '前端可展示 Agent 工具调用时间线和 LangGraph 执行节点；',
    'Gradio 调试台可查看 Agent 执行链、工具调用和检索结果。',
]
doc.add_paragraph('Sprint 2（Agent 与业务融合）：')
for s in sprint2:
    doc.add_paragraph(s, style='List Bullet')

sprint3 = [
    '可通过 Docker Compose 统一启动前后端与数据库；',
    '所有答辩页面风格统一，无明显布局错位或数据加载空白；',
    '即使模型接口临时异常，核心评分、排期、权限模块仍可展示；',
    '测试文档、部署文档、AI 使用记录完整；',
    '三位核心算法负责人能逐行讲解各自禁飞区代码。',
]
doc.add_paragraph('Sprint 3（稳定性、部署与答辩）：')
for s in sprint3:
    doc.add_paragraph(s, style='List Bullet')

doc.add_heading('7.2  需求评审要求', level=2)
doc.add_paragraph(
    '需求评审应重点检查：功能边界是否清晰，角色权限是否完整定义，AI 禁飞区边界是否明确，'
    '薪资权限和审计日志的数据链路是否可追踪，Agent 调用链是否符合"Tool → Service → human_only"约束，'
    '异常流程（LLM 不可用、权限拒绝、RAG 无命中）是否都有明确处理方案。'
)
doc.add_paragraph(
    '评审通过后，后续概要设计、接口契约、数据库设计和测试用例均应与本说明书保持一致。'
    '当实现过程中发现需求不明确或存在冲突时，应先更新需求说明，再同步修改设计、代码和测试用例。'
)

# ── 保存 ──
desktop = os.path.join(os.environ['USERPROFILE'], 'Desktop')
filepath = os.path.join(desktop, 'TalentFlow_软件需求规格说明书.docx')
doc.save(filepath)
print(f'[OK] Document saved to: {filepath}')
